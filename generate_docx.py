from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page setup A4
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Style helpers
def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells_data, bold=False, bg=None, font_color=None):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = "Meiryo"
        run.bold = bold
        if font_color:
            run.font.color.rgb = font_color
        if bg:
            set_cell_shading(cell, bg)
    return row

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Meiryo"
    return h

def para(text, bold=False, color=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Meiryo"
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def red_box(text):
    return para(text, bold=True, color=RGBColor(0xC0, 0x39, 0x2B), size=10)

# ═══════════════════════════════════════
# Title
# ═══════════════════════════════════════
title = doc.add_heading("Fraud Kill Chain（不正送金キルチェーン）", level=0)
for run in title.runs:
    run.font.name = "Meiryo"
para("シート名「不審な電話手口」 ― 犯人サイド × 被害企業サイド 分析レポート", size=11, color=RGBColor(0x66, 0x66, 0x66))
para("対象：法人向けIB（BizSTATION） ｜ 犯人はBizSTATIONの認証フローを完全に把握 ｜ 1件で数千万〜数億円の被害リスク", bold=True, color=RGBColor(0xC0, 0x39, 0x2B), size=10)

doc.add_page_break()

# ═══════════════════════════════════════
# 攻撃の全手順
# ═══════════════════════════════════════
heading("攻撃の全手順（1本の電話の中で完結）")
para("この攻撃は「9つの手順」ではなく、1本の電話の中で起きる1つの連続した攻撃である。手順①から手順⑨まで電話は切れていない。", size=10)

# Step 1
heading("手順① 大量架電（自動音声）", level=2)
para("犯人のシステムが法人に自動音声で一斉架電する。", size=10)
para("「こちらはMバンく BizSTATION法人窓口です。BizSTATIONを利用している方は1を押してください。承認実行権限を持っている方は1を押してください」", size=9, color=RGBColor(0x29, 0x80, 0xB9))
para("• 「1」を押した人 → 手順②へ（権限者と確定）\n• それ以外 → 法人ダイレクトオフィスの番号を案内（通報リスク抑制）", size=9)
para("■ 分析：自動音声で「送金権限を持つ人」だけを自動フィルタリング。「1」以外を押した人にも丁寧に対応し、不審に思われない退出経路を用意。銀行も同じ自動音声振り分けを行うため、顧客には見分けがつかない。", size=9)

# Step 2
heading("手順② 電話で不安を煽る（AI音声 or 有人）", level=2)
para("「パソコン環境の更新が必要。現在口座の利用制限をしています。手続きを郵送していますが、届いてますか。更新されていないので電話しました」", size=9, color=RGBColor(0xE6, 0x7E, 0x22))
para("※AI音声または有人対応（法人高額案件のため両方の可能性あり）", size=9)

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
for i, h in enumerate(["セリフ", "狙い", "効果"]):
    t.rows[0].cells[i].text = h
    set_cell_shading(t.rows[0].cells[i], "222222")
    for p in t.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(8)
            r.font.name = "Meiryo"
            r.bold = True

add_table_row(t, ["「パソコン環境の更新が必要」", "伏線", "後のソフトインストールを自然に受け入れさせる"])
add_table_row(t, ["「口座の利用制限をしています」", "恐怖", "「今すぐ対応しなきゃ」と焦らせ、冷静な判断を奪う"])
add_table_row(t, ["「手続きを郵送しましたが、届いてますか」", "負い目", "「自分が見落とした」と思わせ、指示に従いやすくする"])
add_table_row(t, ["「更新されていないので電話しました」", "正当化", "架電がフォローアップに見え、銀行の丁寧な対応と錯覚させる"])

red_box("🔑 攻撃全体の鍵：「口座の利用制限」で恐怖を植え付け、以降の全手順で判断力を奪う")
para("「届いてますか？」への回答：「届いてない」→「では電話で手続きしましょう」/ 「届いてる」→「確認しながら進めましょう」。どう答えても犯人のシナリオから逸脱できない。", size=9)

# Step 3
heading("手順③ メールアドレスの入手", level=2)
para("「手続きをメールでご案内しますので、メールアドレスを教えてください」", size=9, color=RGBColor(0xE6, 0x7E, 0x22))
red_box("⚠ 破綻ポイント：BizSTATION契約時にメアド登録済み。銀行がメアドを聞く必要はない。= 100%詐欺")
para("→ ただし手順②で「口座制限中」と言われ焦っているため、この矛盾に気づけない。", size=9, color=RGBColor(0x88, 0x88, 0x88))

# Step 4
heading("手順④ フィッシングメール送付", level=2)
para("犯人がメールを送る。電話をつないだまま送っている。", size=10)
para("件名①「【重要】BizSTATIONお客さま情報更新のお知らせ」\n件名②「【Mバンく】インターネットバンキングを安全にご利用いただくための案内（Rapportのご利用推奨）」", size=9)
para("■ 通常のフィッシングとの決定的な違い：電話で「今からメール送ります」と予告してから送る。予告通り届くので「本物だ」と確信してしまう。電話×メールの二重チャネルが相互に信頼性を補強。", size=9)
red_box("⚠ 破綻ポイント：偽メールは正規フォルダに入らない（SPF/DKIM/DMARC不通過）。犯人「迷惑フォルダを確認して」= 銀行が絶対に言わないセリフ = 100%詐欺")
para("→ さらに迷惑フォルダに入っていたことが「重要メールがフィルタに引っかかった→だから口座更新の案内に気づけなかった」と手順②の話を補強してしまう。", size=9, color=RGBColor(0x88, 0x88, 0x88))

# DMARC countermeasure
heading("【技術的対策①】DMARC p=reject ＋ DNS対策", level=3)
para("DMARC p=reject を設定すれば、銀行ドメインを詐称したメールは顧客に一切届かない。手順④でKill Chainが切れる。", size=10, bold=True)

t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Table Grid'
for i, h in enumerate(["DMARCポリシー", "偽メールの扱い", "結果"]):
    t2.rows[0].cells[i].text = h
    set_cell_shading(t2.rows[0].cells[i], "222222")
    for p in t2.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(8)
            r.font.name = "Meiryo"
            r.bold = True
add_table_row(t2, ["p=none（監視のみ）", "普通に届く", "攻撃成立"], bg="FDEDEC")
add_table_row(t2, ["p=quarantine（隔離）", "迷惑フォルダに入る", "電話で突破される"], bg="FDF2E9")
add_table_row(t2, ["p=reject（拒否）★", "メールサーバが受信拒否。配信されない。", "攻撃が止まる"], bg="EAFAF1", font_color=RGBColor(0x27, 0xAE, 0x60))

para("類似ドメインも含めた包括的DNS対策：①類似ドメインの防御的登録 ②新規登録ドメイン（NRD）ブロック ③正規送信ドメインの公開＋GW制御 ④SPF/DKIM未設定ドメインの拒否", size=9)
red_box("技術的にはもう解決できる話。やるかやらないかだけ。")

# Step 5
heading("手順⑤ リモートツールのインストールに誘導", level=2)
para("「セキュリティ強化のためソフトのインストールが必要です。メールのボタンをクリックしてください」", size=9, color=RGBColor(0xE6, 0x7E, 0x22))
para("被害者がクリック → ブラウザにBizSTATIONデザインの画面表示 → setup.msi（ScreenConnect）がダウンロードされる", size=9)
red_box("⚠ 破綻ポイント：金融庁の要請により、金融機関のメールにURLリンクを記載することは禁止。銀行の方針でも「メールのリンクをクリックしないで」と周知済み。= 100%詐欺")
para("■ 皮肉：犯人のほうが「ユーザー体験が良い」。正規の銀行は「ブックマークからアクセスして」（面倒）。犯人は「ボタンをクリックするだけ」（簡単）。", size=9)
para("■ 「ブックマークからアクセスして」は機能しない：ブックマーク登録を顧客に任せているが、登録していない人が多い。銀行がやるべきは、契約時にブックマークを設定してあげる、公式アプリに統一する等。", size=9)

# Step 6
heading("手順⑥ setup.msiの実行", level=2)
para("「セットアップをクリックしてください」→「実行するをクリックしてください」", size=9, color=RGBColor(0xE6, 0x7E, 0x22))
para("msiファイル実行時に「不明な発行元」と表示されるが、犯人の電話指示で突破される。", size=9)
red_box("⚠ 極めつけの証拠：銀行が配布するソフトウェアで「不明な発行元」はありえない。ゼロ。Rapportは「IBM Corporation」、PhishWallは「SecureBrain Corporation」と表示される。")
para("■ ただし、犯人が証明書を購入・窃取すれば「不明な発行元」の警告は消せる。本質的な防壁ではない。", size=9)

# AppLocker countermeasure
heading("【技術的対策②】AppLocker / WDAC（ホワイトリスト型アプリ制御）", level=3)
para("許可されたアプリのみ実行可能にするホワイトリスト型制御。Windowsに標準搭載、追加コストゼロ。", size=10, bold=True)
para("• AppLocker：Windows 10/11 Enterprise / Education に標準搭載\n• WDAC：Windows 10/11 全エディションに標準搭載（AppLockerより強力）\n• 犯人がどんな証明書を持っていても、ホワイトリストにないアプリは実行できない\n• 電話では回避できない（管理者権限がないとポリシー変更できない）", size=9)
red_box("証明書ベースの防御は犯人のレベル次第で突破される。ホワイトリスト型アプリ制御が唯一の確実な対策。")

# Step 7
heading("手順⑦ 赤い画面で目隠し", level=2)
para("犯人はScreenConnect経由で被害者のPCに赤い画面を全画面表示する。赤はMバンくのブランドカラー（コーポレートカラー）に合わせている。", size=10)
para("『Mバンく　現在セキュリティ保護強化に伴うシステム更新ならびに、お客さま情報の確認、更新処理を実施しております。処理完了までの間一時的に画面が遅延または停止する場合がありますが、正常な処理となりますのでそのままお待ちくださいますようお願い申し上げます。』", size=9, color=RGBColor(0xC0, 0x39, 0x2B))

para("■ 赤い画面の3つの役割：\n① PCを「封印」する — 被害者はPCを操作できない。BizSTATIONで自分で確認することも不可能。\n② 犯人の操作を隠す — 赤い画面の裏でScreenConnect経由のPC操作が進行。被害者には見えない。\n③ 被害者を「待機」させる — 「処理完了までお待ちください」で何もせず待たせる。", size=9)

red_box("⚠ 破綻ポイント：銀行のシステム更新で顧客PCの画面が赤くなることはありえない。IBのシステム更新はサーバー側の処理。")
para("■ EDR/NW監視でScreenConnectのプロセス・C2通信を検知可能。ただし中小企業にはEDR/NW監視がない → また「顧客の問題」になる。", size=9)

# Step 8
heading("手順⑧ SMS経由で認証情報を窃取（デュアルデバイス攻撃）", level=2)
para("「手続きを進めるにはSMS認証が必要なため、お客さまの携帯電話の番号を教えてください」", size=9, color=RGBColor(0xE6, 0x7E, 0x22))
red_box("⚠ 破綻ポイント：SMS認証を設定した時点で携帯番号は登録済み。銀行が聞く必要はない。= 100%詐欺")

para("被害者が個人スマホでSMSのリンクをクリック → 偽の認証画面に契約者番号・利用者ID・ログインPW・取引実行PW・OTPの5要素を一括入力させられる。", size=10)
red_box("⚠ 破綻ポイント：正規のBizSTATIONは段階的に認証する。5要素を一画面で入力させること自体がありえない。")

para("■ デュアルデバイス攻撃の構造：\n• 会社PC → 赤い画面で封印済み。確認手段なし。犯人が遠隔操作中。\n• 個人スマホ → SMS経由で認証情報を入力させる。企業セキュリティの管轄外。\n■ OTPの時間制限が犯人に有利：被害者がOTPを入力した瞬間に犯人が即座に使用（有効期限30〜60秒）。電話をつないだままだからリアルタイムにできる。", size=9)

# Step 9
heading("手順⑨ 不正送金の実行", level=2)
para("犯人は赤い画面の裏で被害者のPC（ScreenConnect経由）からBizSTATIONにアクセスし、窃取した認証情報で送金を実行する。", size=10)
para("■ 被害者のPCを経由する理由：犯人のPCからアクセスするとIPアドレスが異なり異常検知される。被害者のPCを経由すれば「いつもの場所からのアクセス」に見える。", size=9)

t3 = doc.add_table(rows=1, cols=4)
t3.style = 'Table Grid'
for i, h in enumerate(["検知項目", "正常", "この攻撃", "結果"]):
    t3.rows[0].cells[i].text = h
    set_cell_shading(t3.rows[0].cells[i], "222222")
    for p in t3.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(8)
            r.font.name = "Meiryo"
            r.bold = True

add_table_row(t3, ["IPアドレス", "いつもの場所", "いつもの場所（PC経由）", "すり抜ける"], bg="FDEDEC")
add_table_row(t3, ["時間帯", "業務時間内", "業務時間内（電話中）", "すり抜ける"], bg="FDEDEC")
add_table_row(t3, ["送金先口座", "取引先", "初めての口座", "検知できる ★"], bg="EAFAF1")
add_table_row(t3, ["送金金額", "通常範囲", "通常より大きい可能性", "検知できる ★"], bg="EAFAF1")
add_table_row(t3, ["操作速度", "人間の速度", "異常に速い可能性", "検知できる ★"], bg="EAFAF1")

para("■ 銀行の追加認証（電話確認）すら機能しない：被害者は犯人との電話がつながったままで銀行からの電話に出られない。出たとしても「はい、自分で送金しました」と答える（攻撃されていることに気づいていない）。", size=9)

doc.add_page_break()

# ═══════════════════════════════════════
# 総括
# ═══════════════════════════════════════
heading("総括：正規の銀行体験を99%再現した攻撃")

heading("1. 正規と同じもの（10個）", level=2)
same = [
    "自動音声の振り分け", "担当者との電話対応", "「郵送した」「口座制限」のセリフ",
    "メールのデザイン（Bizデザイン）", "メールの件名（実在する案内と同じ）",
    "Webページのデザイン", "赤い画面の色（ブランドカラー）",
    "赤い画面の文言（銀行の丁寧語）", "SMS認証の流れ", "セキュリティソフト推奨（Rapportは実在）",
]
for item in same:
    para(f"✅ {item}", size=9, color=RGBColor(0x27, 0xAE, 0x60))

heading("2. 正規と違うもの（6個）— 全て専門知識がないと気づけない", level=2)
diff = [
    ("❌ メアドを聞く", "→ 登録済みなので聞かない"),
    ("❌ 迷惑フォルダを見て", "→ 正規メールは入らない"),
    ("❌ メールにURL", "→ 金融庁が禁止"),
    ("❌ 不明な発行元でも実行して", "→ 正規ソフトは署名済み"),
    ("❌ 携帯番号を聞く", "→ 登録済みなので聞かない"),
    ("❌ 5要素一括入力", "→ 段階的に認証する"),
]
for item, reason in diff:
    para(f"{item}　{reason}", size=9, color=RGBColor(0xC0, 0x39, 0x2B))

heading("3. 「銀行が絶対に言わないセリフ」5つ", level=2)
para("③「メアドを教えて」④「迷惑フォルダを確認して」⑤「メールのボタンをクリックして」⑥「不明な発行元でも実行して」⑧「携帯番号を教えて」", size=9, bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
para("→ 全て手順②の恐怖で無効化される。", size=9)

heading("4. 技術的にKill Chainを切れるポイント — 全て既存ツール・追加コストゼロ", level=2)

t4 = doc.add_table(rows=1, cols=3)
t4.style = 'Table Grid'
for i, h in enumerate(["対策", "効果", "備考"]):
    t4.rows[0].cells[i].text = h
    set_cell_shading(t4.rows[0].cells[i], "27AE60")
    for p in t4.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(8)
            r.font.name = "Meiryo"
            r.bold = True
add_table_row(t4, ["④ DMARC p=reject", "メール遮断（銀行側で完結）", "メールサーバに標準搭載"])
add_table_row(t4, ["⑥ AppLocker/WDAC", "msi実行阻止", "Windowsに標準搭載"])
add_table_row(t4, ["⑦⑨ EDR/振る舞い検知", "ScreenConnect通信＋異常送金検知", "中小企業には支援が必要"])

heading("5. 根本原因と最終結論", level=2)
red_box("全ての対策が「顧客にやってね」で丸投げされ、機能していない。")
para("「メールのリンクをクリックしないで」→ 顧客任せ → 機能しない\n「ブックマークからアクセスして」→ 顧客任せ → 機能しない\n「不明な発行元に注意して」→ 顧客任せ → 機能しない\n「AppLockerを設定してください」→ 顧客任せ → 機能しない", size=9)

para("")
p_final = doc.add_paragraph()
run = p_final.add_run("最終結論：この攻撃が成立している根本原因は、犯人の巧みさではなく「既にある対策を設定していない」「全てを顧客の問題にしている」こと。銀行側が技術的対策を実装すれば、この攻撃は成立しなくなる。")
run.font.size = Pt(11)
run.font.name = "Meiryo"
run.bold = True
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

# ═══════════════════════════════════════
# 銀行の正規フローとの類似性
# ═══════════════════════════════════════
doc.add_page_break()
heading("付録：銀行の正規フローとの類似性")
para("銀行自身が顧客を「自動音声の指示に従う」行動パターンに教育してきた結果、この攻撃が成立している。", size=10)

t5 = doc.add_table(rows=1, cols=3)
t5.style = 'Table Grid'
for i, h in enumerate(["銀行の正規フロー", "この攻撃の手口", "顧客の体験"]):
    t5.rows[0].cells[i].text = h
    set_cell_shading(t5.rows[0].cells[i], "222222")
    for p in t5.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(8)
            r.font.name = "Meiryo"
            r.bold = True

add_table_row(t5, ["自動音声で用件を振り分け", "自動音声で権限者を選別", "同じ"], font_color=RGBColor(0xC0, 0x39, 0x2B))
add_table_row(t5, ["番号を押して選択", "「1」を押す", "同じ"], font_color=RGBColor(0xC0, 0x39, 0x2B))
add_table_row(t5, ["担当者と話す", "犯人（AI音声/有人）と話す", "同じ"], font_color=RGBColor(0xC0, 0x39, 0x2B))
add_table_row(t5, ["指示に従って手続き", "指示に従ってインストール", "同じ"], font_color=RGBColor(0xC0, 0x39, 0x2B))
add_table_row(t5, ["本人確認で情報を伝える", "認証情報をすべて入力", "同じ"], font_color=RGBColor(0xC0, 0x39, 0x2B))

red_box("構造的問題：正規と詐欺の区別が顧客側ではほぼ不可能")

para("\n注意喚起・ガイドライン・金融庁の要請 — すべてが「冷静な時の判断力」を前提にしている。この攻撃は、その前提自体を手順②で破壊している。", size=10, bold=True)

# Save
out = "/home/user/security-news-pwa/Fraud_Kill_Chain.docx"
doc.save(out)
print(f"Saved: {out}")
